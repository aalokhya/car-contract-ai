import os
from docx import Document

REPORT_DIR = "reports"

def generate_report(contract_id: int, analysis: dict):

    if not os.path.exists(REPORT_DIR):
        os.makedirs(REPORT_DIR)

    file_path = os.path.join(REPORT_DIR, f"contract_report_{contract_id}.docx")

    doc = Document()
    doc.add_heading("Contract Analysis Report", level=1)

    # Summary
    doc.add_heading("Contract Summary", level=2)
    for key, value in analysis["summary"].items():
        doc.add_paragraph(f"{key}: {value}")

    # Risks
    doc.add_heading("Risks Detected", level=2)
    for risk in analysis["risks"]:
        doc.add_paragraph(risk)

    # Fairness Score
    doc.add_heading("Fairness Score", level=2)
    doc.add_paragraph(str(analysis["fairness_score"]))

    # Tips
    tips = analysis.get("tips", {})
    if tips:
        doc.add_heading("Negotiation Tips", level=2)

        for item in tips.get("unfair_clauses", []):
            doc.add_paragraph(f"Unfair Clause: {item}")

        for item in tips.get("negotiation_points", []):
            doc.add_paragraph(f"Negotiation Point: {item}")

        doc.add_paragraph("Suggested Message:")
        doc.add_paragraph(tips.get("message_to_dealer", ""))

    doc.save(file_path)

    return file_path
